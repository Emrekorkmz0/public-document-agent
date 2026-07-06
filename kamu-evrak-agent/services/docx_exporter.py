from __future__ import annotations

from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def _safe_text(value: Any, default: str = "") -> str:
    if value is None:
        return default
    text = str(value).strip()
    return text if text else default


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def _set_default_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)


def _add_center_line(doc: Document, text: str, bold: bool = False, size: int = 12) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def _add_label_value(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(11)
    r2 = p.add_run(value)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(11)


def _add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.strip())
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def _split_body(body: str) -> List[str]:
    lines = [line.strip() for line in (body or "").replace("\r\n", "\n").split("\n")]
    paragraphs: List[str] = []
    buf: List[str] = []
    for line in lines:
        if not line:
            if buf:
                paragraphs.append(" ".join(buf).strip())
                buf = []
            continue
        # Madde veya kısa başlık gibi duran satırları tek paragraf tut.
        if line.startswith(("-", "•")):
            if buf:
                paragraphs.append(" ".join(buf).strip())
                buf = []
            paragraphs.append(line)
        else:
            buf.append(line)
    if buf:
        paragraphs.append(" ".join(buf).strip())
    return [p for p in paragraphs if p]


def build_official_docx_bytes(
    analysis: Dict[str, Any],
    routing: Dict[str, Any],
    draft: Dict[str, Any],
    sources: Optional[List[Dict[str, Any]]] = None,
    institution_name: str = "ÖRNEK BELEDİYESİ",
    unit_name: str = "Yazı İşleri Müdürlüğü",
    signer_name: str = "Yetkili Personel",
    signer_title: str = "Birim Yetkilisi",
    document_number: str = "E-00000000-000-000000",
    date_text: Optional[str] = None,
    include_analysis_appendix: bool = True,
) -> bytes:
    """Ekrandaki analiz ve taslak bilgisinden indirilebilir DOCX üretir.

    Bu çıktı nihai resmi belge değil, insan onayı gerektiren ön taslaktır.
    """
    date_text = date_text or datetime.now().strftime("%d.%m.%Y")
    sources = sources or []

    doc = Document()
    _set_default_styles(doc)

    _add_center_line(doc, "T.C.", bold=True, size=12)
    _add_center_line(doc, institution_name.upper(), bold=True, size=12)
    _add_center_line(doc, unit_name, bold=True, size=12)
    doc.add_paragraph()

    # Sayı / konu / tarih alanı
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    values = [
        ("Sayı", document_number),
        ("Konu", _safe_text(draft.get("subject"), "Evrak Hakkında")),
        ("Tarih", date_text),
    ]
    for row, (label, value) in zip(table.rows, values):
        row.cells[0].width = Cm(3.0)
        row.cells[1].width = Cm(12.0)
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_text(row.cells[0], f"{label}:", bold=True)
        _set_cell_text(row.cells[1], value)

    doc.add_paragraph()

    # Muhatap / önerilen birim
    recommended_unit = _safe_text(routing.get("recommended_unit"), "İlgili Birim")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(recommended_unit.upper())
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_paragraph()

    body = _safe_text(draft.get("body"), "İlgili evrakın değerlendirilmesi önerilir.")
    for paragraph in _split_body(body):
        if paragraph.startswith(("-", "•")):
            p = doc.add_paragraph(style=None)
            p.paragraph_format.left_indent = Cm(1.25)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(paragraph)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        else:
            _add_body_paragraph(doc, paragraph)

    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.LEFT
    closing.paragraph_format.first_line_indent = Cm(1.25)
    closing.add_run("Bilgilerinize rica ederim.").font.name = "Times New Roman"

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = sig.add_run(f"{signer_name}\n{signer_title}")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

    source_note = _safe_text(draft.get("source_note"))
    if source_note:
        doc.add_paragraph()
        _add_label_value(doc, "Kaynak Notu: ", source_note)

    _add_label_value(doc, "Dağıtım: ", recommended_unit)

    # İnsan onayı uyarısı
    doc.add_paragraph()
    warn = doc.add_paragraph()
    warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    wr = warn.add_run("Bu belge yapay zekâ destekli ön taslaktır. Nihai kontrol ve onay yetkili personele aittir.")
    wr.italic = True
    wr.font.name = "Times New Roman"
    wr.font.size = Pt(9)

    if include_analysis_appendix:
        doc.add_page_break()
        _add_center_line(doc, "EK: EVRAK ANALİZ NOTU", bold=True, size=12)
        doc.add_paragraph()
        _add_label_value(doc, "Evrak Türü: ", _safe_text(analysis.get("document_type"), "Belirsiz"))
        _add_label_value(doc, "Güven: ", f"%{int(float(analysis.get('confidence', 0.0)) * 100)}")
        _add_label_value(doc, "Risk Seviyesi: ", _safe_text(analysis.get("risk_level"), "-"))
        _add_label_value(doc, "Özet: ", _safe_text(analysis.get("summary"), "-"))
        _add_label_value(doc, "Önerilen Birim: ", recommended_unit)
        _add_label_value(doc, "Yönlendirme Gerekçesi: ", _safe_text(routing.get("reason"), "-"))

        missing = analysis.get("missing_information") or []
        doc.add_paragraph()
        _add_label_value(doc, "Eksik / Belirsiz Bilgiler: ", "Yok" if not missing else "")
        for item in missing:
            p = doc.add_paragraph(style=None)
            p.paragraph_format.left_indent = Cm(0.8)
            run = p.add_run(f"- {item}")
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

        if sources:
            doc.add_paragraph()
            _add_label_value(doc, "RAG Kaynakları: ", "")
            for idx, src in enumerate(sources[:5], start=1):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.8)
                title = _safe_text(src.get("title"), "Kaynak")
                source_type = _safe_text(src.get("source_type"), "-")
                score = src.get("score", 0)
                try:
                    score_text = f"%{int(float(score) * 100)}"
                except Exception:
                    score_text = "-"
                run = p.add_run(f"{idx}. {title} | Tür: {source_type} | Benzerlik: {score_text}")
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
